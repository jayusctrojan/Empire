"""
Empire v7.3 - Metadata Extraction Service
Extracts basic and advanced metadata from various file types
"""

import os
import logging
import hashlib
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional
import mimetypes

# Image metadata
try:
    import exifread
    from PIL import Image
    IMAGE_SUPPORT = True
except ImportError:
    IMAGE_SUPPORT = False
    logging.warning("Image metadata libraries not available (exifread, PIL)")

# Audio/Video metadata
try:
    from mutagen import File as MutagenFile
    AUDIO_VIDEO_SUPPORT = True
except ImportError:
    AUDIO_VIDEO_SUPPORT = False
    logging.warning("Audio/video metadata library not available (mutagen)")

# Document metadata
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.warning("DOCX metadata library not available (python-docx)")

try:
    from pypdf import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logging.warning("PDF metadata library not available (pypdf)")

logger = logging.getLogger(__name__)


class MetadataExtractor:
    """
    Extracts metadata from various file types

    Supports:
    - Images: EXIF data, dimensions, format
    - Audio/Video: Duration, bitrate, codec, artist, album
    - DOCX: Author, title, created/modified dates
    - PDF: Author, title, created/modified dates, page count
    """

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a file

        Args:
            file_path: Path to file

        Returns:
            Dictionary containing metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Extract basic metadata (always available)
        metadata = self._extract_basic_metadata(file_path)

        # Detect file type and extract advanced metadata
        file_ext = Path(file_path).suffix.lower()
        mime_type, _ = mimetypes.guess_type(file_path)

        try:
            # Image files
            if file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                advanced = self._extract_image_metadata(file_path)
                metadata.update(advanced)

            # Audio files
            elif file_ext in ['.mp3', '.wav', '.m4a', '.flac', '.ogg', '.aac']:
                advanced = self._extract_audio_metadata(file_path)
                metadata.update(advanced)

            # Video files
            elif file_ext in ['.mp4', '.mov', '.avi', '.mkv', '.wmv', '.flv']:
                advanced = self._extract_video_metadata(file_path)
                metadata.update(advanced)

            # DOCX files
            elif file_ext == '.docx':
                advanced = self._extract_docx_metadata(file_path)
                metadata.update(advanced)

            # PDF files
            elif file_ext == '.pdf':
                advanced = self._extract_pdf_metadata(file_path)
                metadata.update(advanced)

            # PowerPoint files
            elif file_ext in ['.pptx', '.ppt']:
                advanced = self._extract_pptx_metadata(file_path)
                metadata.update(advanced)

            else:
                logger.debug(f"No advanced metadata extractor for {file_ext}")

        except Exception as e:
            logger.error(f"Error extracting advanced metadata from {file_path}: {e}")
            metadata['metadata_extraction_error'] = str(e)

        return metadata

    def _extract_basic_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract basic file metadata available from filesystem"""
        stat = os.stat(file_path)
        path = Path(file_path)

        # Calculate SHA256 hash of file
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            # Read file in chunks to handle large files efficiently
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        file_hash = sha256_hash.hexdigest()

        return {
            'filename': path.name,
            'file_extension': path.suffix.lower(),
            'file_size_bytes': stat.st_size,
            'file_size_mb': round(stat.st_size / (1024 * 1024), 2),
            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'mime_type': mimetypes.guess_type(file_path)[0],
            'file_hash': file_hash,  # SHA256 hash for integrity verification
        }

    def _extract_image_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract EXIF and image metadata"""
        metadata = {}

        if not IMAGE_SUPPORT:
            return metadata

        try:
            # Get image dimensions using PIL
            with Image.open(file_path) as img:
                metadata['image_width'] = img.width
                metadata['image_height'] = img.height
                metadata['image_format'] = img.format
                metadata['image_mode'] = img.mode  # RGB, RGBA, L, etc.

            # Get EXIF data using exifread
            with open(file_path, 'rb') as f:
                tags = exifread.process_file(f, details=False)

                # Extract common EXIF fields
                exif_data = {}

                # Camera info
                if 'Image Make' in tags:
                    exif_data['camera_make'] = str(tags['Image Make'])
                if 'Image Model' in tags:
                    exif_data['camera_model'] = str(tags['Image Model'])

                # DateTime
                if 'EXIF DateTimeOriginal' in tags:
                    exif_data['datetime_original'] = str(tags['EXIF DateTimeOriginal'])

                # GPS coordinates
                if 'GPS GPSLatitude' in tags and 'GPS GPSLongitude' in tags:
                    exif_data['has_gps'] = True
                    exif_data['gps_latitude'] = str(tags['GPS GPSLatitude'])
                    exif_data['gps_longitude'] = str(tags['GPS GPSLongitude'])
                else:
                    exif_data['has_gps'] = False

                # Exposure settings
                if 'EXIF FNumber' in tags:
                    exif_data['f_number'] = str(tags['EXIF FNumber'])
                if 'EXIF ExposureTime' in tags:
                    exif_data['exposure_time'] = str(tags['EXIF ExposureTime'])
                if 'EXIF ISOSpeedRatings' in tags:
                    exif_data['iso'] = str(tags['EXIF ISOSpeedRatings'])

                metadata['exif_data'] = exif_data

        except Exception as e:
            logger.error(f"Error extracting image metadata: {e}")
            metadata['image_metadata_error'] = str(e)

        return metadata

    def _extract_audio_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract audio metadata using mutagen"""
        metadata = {}

        if not AUDIO_VIDEO_SUPPORT:
            return metadata

        try:
            audio = MutagenFile(file_path)

            if audio is None:
                return metadata

            # Duration
            if hasattr(audio.info, 'length'):
                metadata['duration_seconds'] = round(audio.info.length, 2)
                metadata['duration_formatted'] = self._format_duration(audio.info.length)

            # Bitrate
            if hasattr(audio.info, 'bitrate'):
                metadata['bitrate'] = audio.info.bitrate
                metadata['bitrate_kbps'] = round(audio.info.bitrate / 1000, 2)

            # Sample rate
            if hasattr(audio.info, 'sample_rate'):
                metadata['sample_rate'] = audio.info.sample_rate

            # Channels
            if hasattr(audio.info, 'channels'):
                metadata['channels'] = audio.info.channels

            # Tags (artist, album, title, etc.)
            if audio.tags:
                tags = {}

                # Common tag names across formats
                tag_mapping = {
                    'title': ['TIT2', 'title', '\xa9nam'],
                    'artist': ['TPE1', 'artist', '\xa9ART'],
                    'album': ['TALB', 'album', '\xa9alb'],
                    'album_artist': ['TPE2', 'albumartist', 'aART'],
                    'date': ['TDRC', 'date', '\xa9day'],
                    'genre': ['TCON', 'genre', '\xa9gen'],
                    'track': ['TRCK', 'tracknumber', 'trkn'],
                }

                for key, tag_names in tag_mapping.items():
                    for tag_name in tag_names:
                        if tag_name in audio.tags:
                            tags[key] = str(audio.tags[tag_name][0])
                            break

                if tags:
                    metadata['audio_tags'] = tags

        except Exception as e:
            logger.error(f"Error extracting audio metadata: {e}")
            metadata['audio_metadata_error'] = str(e)

        return metadata

    def _extract_video_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract video metadata using mutagen (similar to audio)"""
        metadata = {}

        if not AUDIO_VIDEO_SUPPORT:
            return metadata

        try:
            video = MutagenFile(file_path)

            if video is None:
                return metadata

            # Duration
            if hasattr(video.info, 'length'):
                metadata['duration_seconds'] = round(video.info.length, 2)
                metadata['duration_formatted'] = self._format_duration(video.info.length)

            # Bitrate
            if hasattr(video.info, 'bitrate'):
                metadata['bitrate'] = video.info.bitrate
                metadata['bitrate_kbps'] = round(video.info.bitrate / 1000, 2)

            # Resolution (if available)
            if hasattr(video.info, 'width') and hasattr(video.info, 'height'):
                metadata['video_width'] = video.info.width
                metadata['video_height'] = video.info.height
                metadata['resolution'] = f"{video.info.width}x{video.info.height}"

            # Frame rate
            if hasattr(video.info, 'fps'):
                metadata['frame_rate'] = video.info.fps

        except Exception as e:
            logger.error(f"Error extracting video metadata: {e}")
            metadata['video_metadata_error'] = str(e)

        return metadata

    def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract DOCX document metadata"""
        metadata = {}

        if not DOCX_SUPPORT:
            return metadata

        try:
            doc = Document(file_path)
            core_props = doc.core_properties

            # Author and title
            if core_props.author:
                metadata['document_author'] = core_props.author
            if core_props.title:
                metadata['document_title'] = core_props.title

            # Dates
            if core_props.created:
                metadata['document_created'] = core_props.created.isoformat()
            if core_props.modified:
                metadata['document_modified'] = core_props.modified.isoformat()

            # Other properties
            if core_props.subject:
                metadata['document_subject'] = core_props.subject
            if core_props.keywords:
                metadata['document_keywords'] = core_props.keywords
            if core_props.comments:
                metadata['document_comments'] = core_props.comments

            # Count paragraphs
            metadata['paragraph_count'] = len(doc.paragraphs)

            # Count tables
            metadata['table_count'] = len(doc.tables)

        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
            metadata['docx_metadata_error'] = str(e)

        return metadata

    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF metadata"""
        metadata = {}

        if not PDF_SUPPORT:
            return metadata

        try:
            reader = PdfReader(file_path)

            # Page count
            metadata['page_count'] = len(reader.pages)

            # PDF info (metadata dictionary)
            if reader.metadata:
                pdf_info = reader.metadata

                # Author
                if pdf_info.author:
                    metadata['document_author'] = pdf_info.author

                # Title
                if pdf_info.title:
                    metadata['document_title'] = pdf_info.title

                # Subject
                if pdf_info.subject:
                    metadata['document_subject'] = pdf_info.subject

                # Creator
                if pdf_info.creator:
                    metadata['pdf_creator'] = pdf_info.creator

                # Producer
                if pdf_info.producer:
                    metadata['pdf_producer'] = pdf_info.producer

                # Creation date
                if pdf_info.creation_date:
                    metadata['pdf_created'] = pdf_info.creation_date.isoformat()

                # Modification date
                if pdf_info.modification_date:
                    metadata['pdf_modified'] = pdf_info.modification_date.isoformat()

        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
            metadata['pdf_metadata_error'] = str(e)

        return metadata

    def _extract_pptx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PowerPoint metadata"""
        metadata = {}

        try:
            from pptx import Presentation

            prs = Presentation(file_path)
            core_props = prs.core_properties

            # Author and title
            if core_props.author:
                metadata['document_author'] = core_props.author
            if core_props.title:
                metadata['document_title'] = core_props.title

            # Dates
            if core_props.created:
                metadata['document_created'] = core_props.created.isoformat()
            if core_props.modified:
                metadata['document_modified'] = core_props.modified.isoformat()

            # Slide count
            metadata['slide_count'] = len(prs.slides)

        except Exception as e:
            logger.error(f"Error extracting PPTX metadata: {e}")
            metadata['pptx_metadata_error'] = str(e)

        return metadata

    def _format_duration(self, seconds: float) -> str:
        """Format duration in seconds to HH:MM:SS"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)

        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"


# Global instance
_metadata_extractor = None


def get_metadata_extractor() -> MetadataExtractor:
    """
    Get singleton instance of MetadataExtractor

    Returns:
        MetadataExtractor instance
    """
    global _metadata_extractor
    if _metadata_extractor is None:
        _metadata_extractor = MetadataExtractor()
    return _metadata_extractor
